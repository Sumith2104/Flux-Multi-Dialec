import sys
try:
    import docx
except ImportError:
    print("Please install python-docx")
    sys.exit(1)

def add_caching_section():
    doc_path = r"c:\Users\sumit\Downloads\Flux-ServerBased-main\Flux-ServerBased-main\Flux-ServerBased-main\fluxbase_ieee_architecture.docx"
    
    try:
        doc = docx.Document(doc_path)
    except Exception as e:
        print(f"Error loading document: {e}")
        return

    # Helper for adding pseudo-headings (bold paragraphs) if native styles are missing
    def add_pseudo_heading(text, size=14):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = docx.shared.Pt(size)

    # Main Section
    doc.add_paragraph() # Spacing
    add_pseudo_heading('VII. Smart Caching and Edge Optimization Architecture', 14)
    
    p = doc.add_paragraph()
    p.add_run('To mitigate the high costs and latency associated with complex analytical NoSQL execution (as described in the SQL Translation layer), Fluxbase employs a multi-tiered smart caching architecture leveraging Next.js Edge caching and Upstash Redis.')

    # Subheading 1
    add_pseudo_heading('A. Global Edge Schema Caching', 12)
    doc.add_paragraph('Database schemas (such as project tables and their associated column data types) are highly volatile during DDL operations but entirely static during heavy read workloads. To optimize reads, Fluxbase utilizes the Vercel Edge Network via the Next.js `unstable_cache` API. Schema definitions are pulled into regional Edge CDN nodes worldwide. When the frontend attempts to map a table layout, the schema is served from memory in <20ms without performing an underlying Firestore read. This cache is instantly manually purged (revalidated via tagging) whenever a strict schema mutation (ADD COLUMN, DROP COLUMN) is committed.')

    # Subheading 2
    add_pseudo_heading('B. Engine-Level Result Memoization (Upstash Redis)', 12)
    doc.add_paragraph('The most computationally expensive actions in external DBaaS clients are analytical queries involving aggregations or `JOIN` commands, which must be polyfilled in memory since the underlying data fabric is a NoSQL document tree. Rather than repeatedly generating these data trees for identical requests (such as a public client dashboard loading on 10,000 devices simultaneously), the Fluxbase `SqlEngine` intercepts heavy `SELECT` queries.')
    
    p2 = doc.add_paragraph()
    p2.add_run('The incoming abstract syntax tree (AST)/query string is hashed utilizing Message-Digest Algorithm 5 (MD5). The engine provisions an HTTP REST call to an ')
    p2.add_run('Upstash Serverless Redis').bold = True
    p2.add_run(' instance to retrieve the hash. If a cache miss occurs, the query runs in full and the resulting JSON object tree is inserted into the cluster with a 60-second Time-To-Live (TTL). Subsequent matching queries immediately retrieve the cached payload, bypassing document reads entirely and protecting the database fabric from throughput spikes.')

    # Subheading 3
    add_pseudo_heading('C. Frontend Query Debouncing', 12)
    doc.add_paragraph('At the client layer, React Query implements a strict 5-minute stale-time protocol coupled with deactivated window-focus refetching loops. This guarantees that user hesitation and rapid interface switching inside the data editor no longer triggers redundant sequential database pulls for historically loaded pagination states.')

    try:
        doc.save(doc_path)
        print("Successfully updated Microsoft Word Architecture document.")
    except Exception as e:
        print(f"Failed to save document: {e}")

if __name__ == "__main__":
    add_caching_section()
